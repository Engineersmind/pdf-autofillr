"""DOCX reader using python-docx."""
from __future__ import annotations
from uploaddocument.readers.base_reader import DocumentReader


class DocxReader(DocumentReader):
    def read(self, file_path: str) -> str:
        from docx import Document
        doc = Document(file_path)
        parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
                if row_text:
                    parts.append(row_text)
        return "\n\n".join(parts)
